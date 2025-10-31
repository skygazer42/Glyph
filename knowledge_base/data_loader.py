"""
数据加载器 - 加载和处理政策文档
"""

import os
import logging
import asyncio
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import json
from datetime import datetime
from uuid import uuid4

from .document_processor import DocumentProcessor
from .enhanced_document_processor import EnhancedDocumentProcessor
from .mineru_adapter import MinerUAdapter
from .vector_store import VectorStore
from ..models.base import PolicyDocument, PolicyType, UUID
from ..config.settings import settings

# 尝试导入文档处理库
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None


class DataLoader:
    """政策文档数据加载器"""

    def __init__(
        self,
        raw_data_dir: str = "resources/data/raw",
        processed_data_dir: str = "resources/data/processed",
        vector_store_path: str = "resources/knowledge_base/vector_store",
        use_enhanced_processor: bool = True
    ):
        self.raw_data_dir = Path(raw_data_dir)
        self.processed_data_dir = Path(processed_data_dir)
        self.vector_store_path = Path(vector_store_path)
        self.logger = logging.getLogger(__name__)
        self.use_enhanced_processor = use_enhanced_processor

        # 确保目录存在
        self.processed_data_dir.mkdir(parents=True, exist_ok=True)
        self.vector_store_path.mkdir(parents=True, exist_ok=True)

        # 初始化组件
        if use_enhanced_processor:
            # 使用增强文档处理器
            config = {
                "mineru_enabled": True,
                "mineru_base_url": settings.model.mineru_base_url,
                "mineru_api_key": settings.model.mineru_api_key,
                "docling_enabled": True,
                "llamaindex_enabled": True,
                "ocr_enabled": settings.document.enable_ocr,
                "table_extraction": settings.document.enable_table_extraction,
                "image_extraction": settings.document.enable_image_extraction
            }
            self.document_processor = EnhancedDocumentProcessor(config)
            self.mineru_adapter = MinerUAdapter()
        else:
            # 使用原始文档处理器
            self.document_processor = DocumentProcessor()
            self.mineru_adapter = None

        self.vector_store = None

    async def load_all_documents(self, force_reload: bool = False) -> List[PolicyDocument]:
        """加载所有政策文档"""
        self.logger.info(f"Loading documents from {self.raw_data_dir}")

        # 检查是否已经处理过
        processed_file = self.processed_data_dir / "processed_documents.json"
        if processed_file.exists() and not force_reload:
            self.logger.info("Loading from processed documents...")
            return await self._load_processed_documents(processed_file)

        # 扫描原始文档
        documents = []
        for file_path in self.raw_data_dir.rglob("*"):
            if file_path.is_file() and self._is_supported_format(file_path):
                try:
                    doc = await self._load_single_document(file_path)
                    if doc:
                        documents.append(doc)
                        self.logger.info(f"Loaded: {file_path.name}")
                except Exception as e:
                    self.logger.error(f"Error loading {file_path}: {e}")

        self.logger.info(f"Total documents loaded: {len(documents)}")

        # 保存处理后的文档
        await self._save_processed_documents(documents, processed_file)

        return documents

    async def _load_single_document(self, file_path: Path) -> Optional[PolicyDocument]:
        """加载单个文档"""
        # 解析文件名获取基本信息
        filename = file_path.stem
        file_content = await self._extract_text_from_file(file_path)

        if not file_content:
            return None

        # 识别政策类型
        policy_type = self._identify_policy_type(filename, file_content)

        # 提取关键信息
        doc_info = await self._extract_document_info(filename, file_content, file_path)

        # 创建PolicyDocument对象
        document = PolicyDocument(
            title=doc_info.get("title", filename),
            content=file_content,
            summary=doc_info.get("summary"),
            source=doc_info.get("source", "未知机构"),
            doc_type=policy_type,
            publish_date=doc_info.get("publish_date"),
            effective_date=doc_info.get("effective_date"),
            expiry_date=doc_info.get("expiry_date"),
            relevant_departments=doc_info.get("departments", []),
            target_groups=doc_info.get("target_groups", []),
            regions=doc_info.get("regions", ["济南市"]),
            keywords=doc_info.get("keywords", []),
            metadata=doc_info.get("metadata", {})
        )

        return document

    async def _extract_text_from_file(self, file_path: Path) -> Optional[str]:
        """从文件中提取文本"""
        if self.use_enhanced_processor:
            # 使用增强文档处理器
            return await self.document_processor.extract_text(file_path)
        else:
            # 使用原始处理方法
            suffix = file_path.suffix.lower()

            try:
                if suffix == ".pdf":
                    return self._extract_from_pdf(file_path)
                elif suffix in [".docx", ".doc"]:
                    return self._extract_from_docx(file_path)
                elif suffix == ".txt":
                    return file_path.read_text(encoding="utf-8")
                elif suffix == ".md":
                    return file_path.read_text(encoding="utf-8")
                else:
                    self.logger.warning(f"Unsupported file format: {suffix}")
                    return None
            except Exception as e:
                self.logger.error(f"Error extracting text from {file_path}: {e}")
                return None

    def _extract_from_pdf(self, file_path: Path) -> str:
        """从PDF提取文本"""
        if PyPDF2 is None:
            self.logger.error("PyPDF2 is not installed. Run: pip install PyPDF2")
            return ""

        try:
            text = []
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
            return "\n".join(text)
        except Exception as e:
            self.logger.error(f"Error extracting PDF {file_path}: {e}")
            return ""

    def _extract_from_docx(self, file_path: Path) -> str:
        """从DOCX提取文本"""
        if docx is None:
            self.logger.error("python-docx is not installed. Run: pip install python-docx")
            return ""

        try:
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return "\n".join(text)
        except Exception as e:
            self.logger.error(f"Error extracting DOCX {file_path}: {e}")
            return ""

    def _identify_policy_type(self, filename: str, content: str) -> PolicyType:
        """识别政策类型"""
        filename_lower = filename.lower()
        content_lower = content.lower()

        # 基于关键词识别
        if any(keyword in filename_lower or keyword in content_lower for keyword in ["以旧换新", "换新"]):
            return PolicyType.REPLACEMENT
        elif any(keyword in filename_lower or keyword in content_lower for keyword in ["消费券", "代金券"]):
            return PolicyType.VOUCHER
        elif any(keyword in filename_lower or keyword in content_lower for keyword in ["补贴", "补助", "津贴"]):
            return PolicyType.SUBSIDY
        elif any(keyword in filename_lower or keyword in content_lower for keyword in ["汽车"]):
            return PolicyType.SUBSIDY  # 汽车补贴归类为补贴
        elif any(keyword in filename_lower or keyword in content_lower for keyword in ["家电"]):
            return PolicyType.REPLACEMENT  # 家电以旧换新
        elif any(keyword in filename_lower or keyword in content_lower for keyword in ["税收", "免税", "减税"]):
            return PolicyType.TAX_EXEMPTION
        elif any(keyword in filename_lower or keyword in content_lower for keyword in ["规定", "办法", "条例"]):
            return PolicyType.REGULATION
        else:
            return PolicyType.GUIDELINE

    async def _extract_document_info(
        self,
        filename: str,
        content: str,
        file_path: Path
    ) -> Dict[str, Any]:
        """提取文档信息"""
        info = {}

        # 提取标题（优先使用文件名）
        info["title"] = self._clean_title(filename)

        # 提取发布机构
        info["source"] = self._extract_source(content)

        # 提取日期
        info["publish_date"] = self._extract_date(content)
        info["effective_date"] = self._extract_date(content, terms=["生效", "实施"])
        info["expiry_date"] = self._extract_date(content, terms=["截止", "失效", "有效期至"])

        # 提取部门
        info["departments"] = self._extract_departments(content)

        # 提取目标群体
        info["target_groups"] = self._extract_target_groups(content)

        # 提取地区
        info["regions"] = self._extract_regions(content)

        # 生成摘要
        info["summary"] = self._generate_summary(content)

        # 提取关键词
        info["keywords"] = self._extract_keywords(content, filename)

        # 元数据
        info["metadata"] = {
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "processed_at": datetime.now().isoformat()
        }

        return info

    def _clean_title(self, filename: str) -> str:
        """清理标题"""
        # 移除常见的文件名前缀和后缀
        prefixes = ["关于印发", "关于", "（", "(", "】", "]", "1.", "2.", "3."]
        suffixes = [".pdf", ".docx", ".doc", ".txt", ".md", "的通知", "的公告", "(1)", "(2)", "(3)"]

        title = filename
        for prefix in prefixes:
            if title.startswith(prefix):
                title = title[len(prefix):]
        for suffix in suffixes:
            if title.endswith(suffix):
                title = title[:-len(suffix)]

        return title.strip()

    def _extract_source(self, content: str) -> str:
        """提取发布机构"""
        # 常见机构模式
        patterns = [
            r"([^。\n]*(?:商务厅|财政局|发改委|工信局|人社局|人民政府)[^。\n]*)",
            r"发布单位[：:]\s*([^\n]+)",
            r"发文单位[：:]\s*([^\n]+)"
        ]

        for pattern in patterns:
            match = re.search(pattern, content)
            if match:
                return match.group(1).strip()

        return "济南市相关部门"

    def _extract_date(self, content: str, terms: List[str] = None) -> Optional[datetime]:
        """提取日期"""
        if terms is None:
            terms = ["发布", "印发", "发文"]

        # 日期模式
        patterns = []
        for term in terms:
            patterns.extend([
                rf"{term}日期[：:]?\s*(\d{{4}}年\d{{1,2}}月\d{{1,2}}日)",
                rf"{term}[：:]?\s*(\d{{4}}年\d{{1,2}}月\d{{1,2}}日)",
                rf"(\d{{4}}年\d{{1,2}}月\d{{1,2}}日){term}"
            ])

        # 添加通用日期模式
        patterns.extend([
            r"(\d{4}年\d{1,2}月\d{1,2}日)",
            r"(\d{4}-\d{1,2}-\d{1,2})"
        ])

        for pattern in patterns:
            match = re.search(pattern, content)
            if match:
                date_str = match.group(1)
                try:
                    # 转换为datetime对象
                    date_str = date_str.replace("年", "-").replace("月", "-").replace("日", "")
                    return datetime.strptime(date_str, "%Y-%m-%d")
                except:
                    continue

        return None

    def _extract_departments(self, content: str) -> List[str]:
        """提取相关部门"""
        departments = []

        # 部门关键词
        dept_keywords = [
            "商务厅", "财政局", "发改委", "工信局", "人社局",
            "税务局", "市场监督管理局", "住房和城乡建设局",
            "交通运输局", "教育局", "卫生健康委"
        ]

        for keyword in dept_keywords:
            if keyword in content:
                departments.append(keyword)

        return list(set(departments))

    def _extract_target_groups(self, content: str) -> List[str]:
        """提取目标群体"""
        groups = []

        # 群体关键词
        group_patterns = [
            r"(城乡居民|居民|市民)",
            r"(企业|公司|商户|个体户)",
            r"(高校毕业生|大学生|毕业生)",
            r"(退役军人|军人)",
            r"(残疾人|残障人士)",
            r"(老年人|老人|退休人员)",
            r"(低收入家庭|困难家庭)",
            r"(小微企业|中小微企业)"
        ]

        for pattern in group_patterns:
            matches = re.findall(pattern, content)
            groups.extend(matches)

        return list(set(groups))

    def _extract_regions(self, content: str) -> List[str]:
        """提取适用地区"""
        regions = []

        # 地区关键词
        if "全省" in content:
            regions.append("全省")
        if "济南市" in content or "济南" in content:
            regions.append("济南市")

        # 提取具体区县
        district_pattern = r"(历下区|市中区|槐荫区|天桥区|历城区|长清区|章丘区|济阳区|莱芜区|钢城区|平阴县|商河县)"
        matches = re.findall(district_pattern, content)
        regions.extend(matches)

        return list(set(regions)) if regions else ["济南市"]

    def _generate_summary(self, content: str, max_length: int = 200) -> str:
        """生成摘要"""
        # 简单提取前几段作为摘要
        sentences = content.split("。")
        summary = ""
        for sentence in sentences[:3]:
            if len(summary) + len(sentence) < max_length:
                summary += sentence + "。"
            else:
                break

        return summary.strip() if summary else content[:max_length] + "..."

    def _extract_keywords(self, content: str, filename: str) -> List[str]:
        """提取关键词"""
        keywords = []

        # 政策相关关键词
        policy_keywords = [
            "补贴", "申请", "条件", "流程", "材料", "时间", "期限",
            "标准", "金额", "资格", "要求", "办法", "实施细则",
            "以旧换新", "消费券", "汽车", "家电", "新能源汽车"
        ]

        content_lower = content.lower()
        filename_lower = filename.lower()

        for keyword in policy_keywords:
            if keyword in content_lower or keyword in filename_lower:
                keywords.append(keyword)

        return list(set(keywords))

    def _is_supported_format(self, file_path: Path) -> bool:
        """检查是否支持的文件格式"""
        supported_formats = [".pdf", ".docx", ".doc", ".txt", ".md"]
        return file_path.suffix.lower() in supported_formats

    async def _load_processed_documents(self, processed_file: Path) -> List[PolicyDocument]:
        """加载已处理的文档"""
        with open(processed_file, "r", encoding="utf-8") as f:
            data = json.load(f)

        documents = []
        for doc_data in data:
            document = PolicyDocument(**doc_data)
            documents.append(document)

        return documents

    async def _save_processed_documents(
        self,
        documents: List[PolicyDocument],
        processed_file: Path
    ):
        """保存处理后的文档"""
        # 转换为可序列化的格式
        data = []
        for doc in documents:
            doc_dict = doc.dict()
            # 处理UUID和datetime
            doc_dict["id"] = str(doc_dict["id"])
            if doc_dict.get("publish_date"):
                doc_dict["publish_date"] = doc_dict["publish_date"].isoformat()
            if doc_dict.get("effective_date"):
                doc_dict["effective_date"] = doc_dict["effective_date"].isoformat()
            if doc_dict.get("expiry_date"):
                doc_dict["expiry_date"] = doc_dict["expiry_date"].isoformat()
            data.append(doc_dict)

        # 保存到文件
        with open(processed_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        self.logger.info(f"Saved {len(documents)} processed documents to {processed_file}")

    async def initialize_vector_store(self, documents: List[PolicyDocument]):
        """初始化向量存储"""
        self.logger.info("Initializing vector store...")

        # 初始化向量存储
        self.vector_store = VectorStore(
            model_name="BAAI/bge-large-zh-v1.5",
            index_path=str(self.vector_store_path / "policy_index.faiss"),
            metadata_path=str(self.vector_store_path / "metadata.pkl")
        )

        # 添加文档到向量存储
        await self.vector_store.add_documents(documents)

        self.logger.info(f"Vector store initialized with {len(documents)} documents")

    async def load_and_initialize(self, force_reload: bool = False):
        """加载文档并初始化向量存储"""
        # 加载文档
        documents = await self.load_all_documents(force_reload=force_reload)

        if not documents:
            self.logger.warning("No documents loaded!")
            return

        # 初始化向量存储
        await self.initialize_vector_store(documents)

        return documents