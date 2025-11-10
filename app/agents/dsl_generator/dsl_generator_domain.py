"""
领域驱动的DSL生成器 - 核心实现
Domain-Driven DSL Generator with Schema Constraints
"""

import os
import json
import yaml
import docx
import jinja2
import re
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from pydantic import ValidationError

# 导入配置和模式
from app.domains import domain_config, SCHEMA_REGISTRY

logger = logging.getLogger(__name__)


class DomainDrivenDSLGenerator:
    """
    领域驱动的DSL生成器
    - 自动识别政策领域
    - 使用专用提示词和模板
    - Schema约束验证
    - 术语标准化
    """

    def __init__(self,
                 api_key: Optional[str] = None,
                 api_base: Optional[str] = None,
                 output_dir: str = "rules"):
        """
        初始化生成器

        Args:
            api_key: LLM API密钥
            api_base: LLM API基础URL
            output_dir: DSL文件输出目录
        """
        self.api_key = api_key or os.getenv("LLM_API_KEY")
        self.api_base = api_base or os.getenv("LLM_BASE_URL")
        self.model = os.getenv("LLM_MODEL_NAME", "qwen-turbo")
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

        # 初始化Jinja2环境
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(['.', 'templates']),
            autoescape=False
        )

        logger.info(f"DSL生成器初始化完成，输出目录: {self.output_dir}")

    def extract_from_docx(self, docx_path: Union[str, Path]) -> str:
        """
        从Word文档提取文本

        Args:
            docx_path: Word文档路径

        Returns:
            提取的文本内容
        """
        doc = docx.Document(docx_path)
        paragraphs = []

        # 提取段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n".join(paragraphs)

    def detect_domain(self, text: str) -> Optional[str]:
        """
        自动检测政策领域

        Args:
            text: 政策文本

        Returns:
            检测到的领域名称
        """
        domain = domain_config.detect_domain(text)
        if domain:
            logger.info(f"检测到政策领域: {domain}")
        else:
            logger.warning("无法自动识别政策领域")
        return domain

    def extract_with_llm(self, text: str, domain: str) -> Dict[str, Any]:
        """
        使用LLM提取结构化数据

        Args:
            text: 政策文本
            domain: 政策领域

        Returns:
            提取的结构化数据
        """
        # 获取领域专用提示词
        prompt_template = domain_config.get_prompt_template(domain)
        if not prompt_template:
            raise ValueError(f"未找到领域 {domain} 的提示词模板")

        # 渲染提示词
        prompt = prompt_template.replace("{{policy_text}}", text)

        # 调用LLM
        try:
            from openai import OpenAI

            client = OpenAI(
                api_key=self.api_key,
                base_url=self.api_base
            )

            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": f"你是{domain}政策规则提取专家"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0,  # 降低随机性
                max_tokens=2000
            )

            content = response.choices[0].message.content
            logger.debug(f"LLM响应: {content[:500]}...")

            # 解析JSON响应
            return self._parse_json_response(content)

        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            raise

    def _parse_json_response(self, response: str) -> Dict[str, Any]:
        """解析LLM返回的JSON"""
        try:
            # 直接尝试解析
            return json.loads(response)
        except json.JSONDecodeError:
            # 尝试提取JSON部分
            json_pattern = r'\{[\s\S]*\}'
            json_match = re.search(json_pattern, response)
            if json_match:
                try:
                    return json.loads(json_match.group())
                except:
                    pass

            # 尝试修复常见问题
            fixed = response
            # 修复尾随逗号
            fixed = re.sub(r',(\s*[}\]])', r'\1', fixed)
            # 修复单引号
            fixed = fixed.replace("'", '"')

            try:
                return json.loads(fixed)
            except:
                logger.error("无法解析LLM响应为JSON")
                raise ValueError("LLM响应格式错误")

    def validate_with_schema(self, data: Dict[str, Any], domain: str) -> Dict[str, Any]:
        """
        使用Schema验证和规范化数据

        Args:
            data: 待验证的数据
            domain: 政策领域

        Returns:
            验证后的数据
        """
        schema_class = SCHEMA_REGISTRY.get(domain)
        if not schema_class:
            logger.warning(f"未找到领域 {domain} 的Schema，跳过验证")
            return data

        try:
            # 使用Pydantic Schema验证
            schema_instance = schema_class(**data)
            validated_data = schema_instance.dict()
            logger.info("Schema验证通过")
            return validated_data

        except ValidationError as e:
            logger.warning(f"Schema验证失败: {e}")
            # 尝试修复常见问题
            return self._fix_common_issues(data, domain)

    def _fix_common_issues(self, data: Dict[str, Any], domain: str) -> Dict[str, Any]:
        """修复常见的数据问题"""
        fixed = data.copy()

        # 修复rule_id格式
        if 'rule_id' in fixed:
            if domain == "汽车补贴" and "Car" not in fixed['rule_id']:
                fixed['rule_id'] = re.sub(r'_\w+_(\d{4})$', f'_Car_\\1', fixed['rule_id'])
            elif domain == "家电补贴" and "Appliance" not in fixed['rule_id']:
                fixed['rule_id'] = re.sub(r'_\w+_(\d{4})$', f'_Appliance_\\1', fixed['rule_id'])
            elif domain == "消费券" and "Coupon" not in fixed['rule_id']:
                fixed['rule_id'] = re.sub(r'_\w+_(\d{4})$', f'_Coupon_\\1', fixed['rule_id'])

        # 修复时间格式（24:00 -> 次日00:00）
        def fix_time(time_str):
            if isinstance(time_str, str) and "24:00" in time_str:
                # 提取日期部分
                date_match = re.match(r'(\d{4}-\d{2}-\d{2})', time_str)
                if date_match:
                    from datetime import datetime, timedelta
                    date = datetime.strptime(date_match.group(1), "%Y-%m-%d")
                    next_day = date + timedelta(days=1)
                    return next_day.strftime("%Y-%m-%d 00:00:00")
            return time_str

        # 递归修复时间字段
        def fix_times_recursive(obj):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if 'time' in key.lower() or 'start' in key or 'end' in key:
                        obj[key] = fix_time(value)
                    elif isinstance(value, (dict, list)):
                        fix_times_recursive(value)
            elif isinstance(obj, list):
                for item in obj:
                    fix_times_recursive(item)

        fix_times_recursive(fixed)

        # 标准化术语
        if domain in domain_config.domains:
            normalization = domain_config.domains[domain].get("normalization", {})
            # 应用术语标准化
            def normalize_recursive(obj):
                if isinstance(obj, str):
                    for standard, variants in normalization.items():
                        if isinstance(variants, list):
                            for variant in variants:
                                if variant in obj:
                                    return obj.replace(variant, standard)
                    return obj
                elif isinstance(obj, dict):
                    return {k: normalize_recursive(v) for k, v in obj.items()}
                elif isinstance(obj, list):
                    return [normalize_recursive(item) for item in obj]
                return obj

            fixed = normalize_recursive(fixed)

        return fixed

    def generate_dsl_yaml(self, data: Dict[str, Any], domain: str) -> str:
        """
        生成DSL YAML

        Args:
            data: 结构化数据
            domain: 政策领域

        Returns:
            生成的YAML字符串
        """
        # 获取领域专用模板
        template_path = domain_config.domains[domain]["dsl_template"]
        template = self.jinja_env.get_template(template_path)

        # 渲染模板
        yaml_content = template.render(**data)

        logger.info(f"成功生成 {domain} DSL YAML")
        return yaml_content

    def save_dsl(self, yaml_content: str, rule_id: str) -> Path:
        """
        保存DSL到文件

        Args:
            yaml_content: YAML内容
            rule_id: 规则ID

        Returns:
            保存的文件路径
        """
        file_path = self.output_dir / f"{rule_id}.yaml"
        file_path.write_text(yaml_content, encoding='utf-8')
        logger.info(f"DSL已保存到: {file_path}")
        return file_path

    def generate_from_text(self, text: str, domain: Optional[str] = None) -> Path:
        """
        从文本生成DSL（完整流程）

        Args:
            text: 政策文本
            domain: 政策领域（可选，会自动检测）

        Returns:
            生成的DSL文件路径
        """
        # 1. 检测领域
        if not domain:
            domain = self.detect_domain(text)
            if not domain:
                raise ValueError("无法识别政策领域，请手动指定domain参数")

        logger.info(f"处理 {domain} 政策")

        # 2. LLM提取
        extracted_data = self.extract_with_llm(text, domain)

        # 3. Schema验证
        validated_data = self.validate_with_schema(extracted_data, domain)

        # 4. 生成YAML
        yaml_content = self.generate_dsl_yaml(validated_data, domain)

        # 5. 保存文件
        rule_id = validated_data.get('rule_id', f'Rule_Unknown_{datetime.now().strftime("%Y%m%d")}')
        file_path = self.save_dsl(yaml_content, rule_id)

        return file_path

    def generate_from_docx(self, docx_path: Union[str, Path], domain: Optional[str] = None) -> Path:
        """
        从Word文档生成DSL

        Args:
            docx_path: Word文档路径
            domain: 政策领域（可选）

        Returns:
            生成的DSL文件路径
        """
        # 提取文档文本
        text = self.extract_from_docx(docx_path)

        # 生成DSL
        return self.generate_from_text(text, domain)

    def batch_generate(self, input_dir: Union[str, Path], domain_hints: Optional[Dict[str, str]] = None):
        """
        批量生成DSL

        Args:
            input_dir: 输入文档目录
            domain_hints: 文件名到领域的映射提示

        Returns:
            生成的DSL文件列表
        """
        input_path = Path(input_dir)
        results = []
        domain_hints = domain_hints or {}

        for docx_file in input_path.glob("*.docx"):
            try:
                logger.info(f"处理文档: {docx_file.name}")
                domain = domain_hints.get(docx_file.stem)
                dsl_path = self.generate_from_docx(docx_file, domain)
                results.append({
                    "source": docx_file.name,
                    "dsl": dsl_path.name,
                    "status": "success"
                })
            except Exception as e:
                logger.error(f"处理 {docx_file.name} 失败: {e}")
                results.append({
                    "source": docx_file.name,
                    "error": str(e),
                    "status": "failed"
                })

        # 生成报告
        self._generate_report(results)
        return results

    def _generate_report(self, results: List[Dict[str, Any]]):
        """生成处理报告"""
        report_path = self.output_dir / "generation_report.json"
        report = {
            "timestamp": datetime.now().isoformat(),
            "total": len(results),
            "success": sum(1 for r in results if r["status"] == "success"),
            "failed": sum(1 for r in results if r["status"] == "failed"),
            "details": results
        }
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2))
        logger.info(f"生成报告已保存: {report_path}")


# 便捷函数
def generate_dsl(text_or_path: Union[str, Path], domain: Optional[str] = None) -> Path:
    """
    便捷函数：生成DSL

    Args:
        text_or_path: 政策文本或文档路径
        domain: 政策领域

    Returns:
        生成的DSL文件路径
    """
    generator = DomainDrivenDSLGenerator()

    if isinstance(text_or_path, Path) or (isinstance(text_or_path, str) and text_or_path.endswith('.docx')):
        return generator.generate_from_docx(text_or_path, domain)
    else:
        return generator.generate_from_text(text_or_path, domain)


if __name__ == "__main__":
    # 测试示例
    import sys

    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        domain = sys.argv[2] if len(sys.argv) > 2 else None

        try:
            dsl_path = generate_dsl(input_file, domain)
            print(f"✅ DSL生成成功: {dsl_path}")
        except Exception as e:
            print(f"❌ DSL生成失败: {e}")
    else:
        print("用法: python dsl_generator_domain.py <输入文件> [领域]")
        print("领域选项: 汽车补贴, 家电补贴, 消费券, 以旧换新, 保险补贴")