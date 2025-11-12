"""
Document loading utilities for the policy QA system.
"""

import os
import uuid
from typing import List, Optional, Dict, Any
from pathlib import Path
from datetime import datetime
import PyPDF2
import re

try:  # python-docx 可能未安装
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

from app.models.base import PolicyDocument, PolicyType


class DocumentLoader:
    """Load and process policy documents from various sources."""

    def __init__(self):
        """Initialize the document loader."""
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md', '.doc'}

    def iter_documents_from_directory(
        self,
        directory_path: str,
        *,
        limit: Optional[int] = None,
    ):
        """Yield documents from a directory to avoid loading everything into memory."""
        directory = Path(directory_path)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        yielded = 0
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.load_single_file(str(file_path))
                    if doc:
                        yield doc
                        yielded += 1
                        if limit is not None and yielded >= limit:
                            break
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")

    def load_from_directory(self, directory_path: str, *, limit: Optional[int] = None) -> List[PolicyDocument]:
        """Load supported documents from a directory (kept for compatibility)."""
        return list(self.iter_documents_from_directory(directory_path, limit=limit))

    def load_single_file(self, file_path: str) -> Optional[PolicyDocument]:
        """Load a single document file."""
        file_path = Path(file_path)

        if not file_path.exists():
            return None

        # Extract content based on file type
        content = self._extract_content(file_path)
        if not content:
            return None

        # Extract metadata
        metadata = self._extract_metadata(file_path, content)

        # Create policy document
        return PolicyDocument(
            id=uuid.uuid4(),
            title=metadata.get('title', file_path.stem),
            content=content,
            summary=self._generate_summary(content),
            source=metadata.get('source', self._infer_source(file_path)),
            doc_type=metadata.get('doc_type', self._infer_policy_type(content)),
            publish_date=metadata.get('publish_date'),
            effective_date=metadata.get('effective_date'),
            expiry_date=metadata.get('expiry_date'),
            relevant_departments=metadata.get('departments', []),
            target_groups=metadata.get('target_groups', []),
            regions=metadata.get('regions', [self._infer_region(file_path)]),
            keywords=metadata.get('keywords', self._extract_keywords(content)),
            metadata=metadata
        )

    def _extract_content(self, file_path: Path) -> Optional[str]:
        """Extract text content from file."""
        extension = file_path.suffix.lower()

        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension == '.docx':
                if Document is None:
                    print("python-docx 未安装，跳过 DOCX 文件解析")
                    return None
                return self._extract_from_docx(file_path)
            elif extension in ['.txt', '.md']:
                return self._extract_from_text(file_path)
            else:
                return None
        except Exception as e:
            print(f"Error extracting content from {file_path}: {e}")
            return None

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        content = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    content.append(text)

        return '\n'.join(content)

    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        content = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(' | '.join(row_text))

        return '\n'.join(content)

    def _extract_from_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract metadata from file path and content."""
        metadata = {}

        # Extract title
        metadata['title'] = self._extract_title(content) or file_path.stem

        # Extract dates
        metadata.update(self._extract_dates(content))

        # Extract departments
        metadata['departments'] = self._extract_departments(content)

        # Extract target groups
        metadata['target_groups'] = self._extract_target_groups(content)

        # Extract regions
        metadata['regions'] = self._extract_regions(content)

        return metadata

    def _extract_title(self, content: str) -> Optional[str]:
        """Extract title from content."""
        lines = content.split('\n')[:10]  # Check first 10 lines

        for line in lines:
            line = line.strip()
            # Look for title patterns
            if (len(line) > 10 and len(line) < 100 and
                any(keyword in line for keyword in ['关于', '通知', '办法', '方案', '细则', '规定', '意见'])):
                return line

        return None

    def _extract_dates(self, content: str) -> Dict[str, Optional[datetime]]:
        """Extract dates from content."""
        dates = {}

        # Date patterns
        patterns = [
            (r'(\d{4}年\d{1,2}月\d{1,2}日)', 'publish_date'),
            (r'(\d{4}年\d{1,2}月)', 'publish_date'),
            (r'自(\d{4}年\d{1,2}月\d{1,2}日)起', 'effective_date'),
            (r'有效期至(\d{4}年\d{1,2}月\d{1,2}日)', 'expiry_date'),
            (r'截止(\d{4}年\d{1,2}月\d{1,2}日)', 'expiry_date')
        ]

        content_preview = content[:1000]  # Check first 1000 characters

        for pattern, date_type in patterns:
            match = re.search(pattern, content_preview)
            if match:
                try:
                    date_str = match.group(1)
                    date = self._parse_chinese_date(date_str)
                    dates[date_type] = date
                except:
                    pass

        return dates

    def _parse_chinese_date(self, date_str: str) -> datetime:
        """Parse Chinese date string."""
        # Simple parsing for Chinese dates
        patterns = [
            (r'(\d{4})年(\d{1,2})月(\d{1,2})日', '%Y-%m-%d'),
            (r'(\d{4})年(\d{1,2})月', '%Y-%m-01')
        ]

        for pattern, format_str in patterns:
            match = re.search(pattern, date_str)
            if match:
                if len(match.groups()) == 3:
                    year, month, day = match.groups()
                    return datetime(int(year), int(month), int(day))
                else:
                    year, month = match.groups()
                    return datetime(int(year), int(month), 1)

        return datetime.now()

    def _extract_departments(self, content: str) -> List[str]:
        """Extract government departments from content."""
        departments = []
        patterns = [
            r'(山东省[\u4e00-\u9fff]+[厅局委部])',
            r'(济南市[\u4e00-\u9fff]+[厅局委部])',
            r'([\u4e00-\u9fff]+市[\u4e00-\u9fff]+[厅局委部])',
            r'([\u4e00-\u9fff]+部门)',
            r'([\u4e00-\u9fff]+办公室)'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, content)
            departments.extend(matches)

        return list(set(departments))  # Remove duplicates

    def _extract_target_groups(self, content: str) -> List[str]:
        """Extract target groups from content."""
        groups = []
        keywords = [
            '个人', '企业', '居民', '农户', '个体工商户',
            '小微企业', '国有企业', '民营企业',
            '老年人', '残疾人', '学生', '失业人员',
            '低收入家庭', '困难群体'
        ]

        for keyword in keywords:
            if keyword in content:
                groups.append(keyword)

        return list(set(groups))

    def _extract_regions(self, content: str) -> List[str]:
        """Extract regions from content."""
        regions = []
        patterns = [
            r'(山东省)',
            r'(济南市)',
            r'([青烟威][岛海坊])',  # Qingdao, Yantai, Weihai, Weifang
            r'([京津沪渝])',  # Municipalities
            r'([\u4e00-\u9fff]+[省市自治区])'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, content)
            regions.extend(matches)

        return list(set(regions))

    def _extract_keywords(self, content: str) -> List[str]:
        """Extract keywords from content."""
        keywords = []
        policy_keywords = [
            '补贴', '补助', '津贴', '奖励',
            '税收', '减免', '优惠', '返还',
            '消费券', '代金券', '优惠券',
            '以旧换新', '更新换代',
            '申请', '审批', '审核', '公示',
            '资金', '预算', '财政',
            '新能源汽车', '家电', '汽车',
            '住房', '租房', '购房',
            '创业', '就业', '培训',
            '社保', '医保', '养老'
        ]

        for keyword in policy_keywords:
            if keyword in content:
                keywords.append(keyword)

        return list(set(keywords))

    def _generate_summary(self, content: str) -> str:
        """Generate a summary of the content."""
        # Simple summary: first paragraph or first 200 characters
        paragraphs = content.split('\n\n')
        if paragraphs:
            first_para = paragraphs[0].strip()
            if len(first_para) > 50:
                return first_para[:500] + '...' if len(first_para) > 500 else first_para

        return content[:200] + '...' if len(content) > 200 else content

    def _infer_source(self, file_path: Path) -> str:
        """Infer source from file path."""
        path_str = str(file_path).lower()

        if '山东' in path_str:
            return '山东省政府'
        elif '济南' in path_str:
            return '济南市政府'
        elif '省' in path_str:
            return '省级部门'
        elif '市' in path_str:
            return '市级部门'
        else:
            return '未知来源'

    def _infer_policy_type(self, content: str) -> PolicyType:
        """Infer policy type from content."""
        content_lower = content.lower()

        if '以旧换新' in content_lower:
            return PolicyType.REPLACEMENT
        elif '消费券' in content_lower or '代金券' in content_lower:
            return PolicyType.VOUCHER
        elif '补贴' in content_lower or '补助' in content_lower:
            return PolicyType.SUBSIDY
        elif '税收' in content_lower or '减免' in content_lower:
            return PolicyType.TAX_EXEMPTION
        elif '办法' in content_lower or '规定' in content_lower:
            return PolicyType.REGULATION
        else:
            return PolicyType.GUIDELINE

    def _infer_region(self, file_path: Path) -> str:
        """Infer region from file path."""
        path_str = str(file_path)

        if '山东' in path_str:
            return '山东省'
        elif '济南' in path_str:
            return '济南市'
        else:
            return '未知地区'
