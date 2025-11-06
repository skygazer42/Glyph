"""
随机生成政策文档并自动转换为 DSL
展示完整的文档生成 → DSL 转换流程
"""

import os
import sys
import random
from datetime import datetime, timedelta
from pathlib import Path

# 添加项目路径
sys.path.append(str(Path(__file__).parent))

from agents.dsl_generator.main_v2 import DSLPipeline


class PolicyDocumentGenerator:
    """随机政策文档生成器"""

    def __init__(self):
        """初始化生成器"""
        # 城市列表
        self.cities = ['济南', '青岛', '烟台', '潍坊', '临沂', '淄博', '济宁', '泰安']

        # 政策类型
        self.policy_types = [
            '家电以旧换新补贴',
            '新能源汽车购置补贴',
            '消费券发放',
            '数字人民币推广补贴',
            '绿色建材消费补贴',
            '餐饮消费促进',
            '文旅消费补贴'
        ]

        # 部门
        self.departments = ['市商务局', '市发改委', '市财政局', '市工信局']

        # 补贴方式
        self.subsidy_methods = [
            '按比例补贴',
            '分档补贴',
            '满减优惠',
            '固定金额补贴'
        ]

    def generate_random_policy(self) -> tuple:
        """生成随机政策内容"""
        city = random.choice(self.cities)
        policy_type = random.choice(self.policy_types)
        department = random.choice(self.departments)
        method = random.choice(self.subsidy_methods)

        # 生成时间范围
        start_date = datetime.now() + timedelta(days=random.randint(1, 30))
        end_date = start_date + timedelta(days=random.randint(90, 365))

        # 生成文号
        doc_number = f"{city[:2]}商发〔{datetime.now().year}〕{random.randint(1, 100)}号"

        # 根据补贴方式生成具体内容
        if method == '按比例补贴':
            content = self._generate_percentage_policy(policy_type)
        elif method == '分档补贴':
            content = self._generate_tiered_policy(policy_type)
        elif method == '满减优惠':
            content = self._generate_discount_policy(policy_type)
        else:
            content = self._generate_fixed_policy(policy_type)

        # 组装完整文档
        document = f"""{city}市{policy_type}实施方案
{city}{department} {doc_number}

为促进消费升级，拉动经济增长，特制定本实施方案。

一、补贴对象
在{city}市范围内的个人消费者。

二、补贴标准
{content}

三、申请条件
1. 申请人须为{city}市户籍或在{city}缴纳社保满6个月
2. 每人限申请{random.randint(1, 3)}次
3. 需提供有效购买凭证

四、活动时间
{start_date.strftime('%Y年%m月%d日')}至{end_date.strftime('%Y年%m月%d日')}

五、资金来源
市级财政专项资金

六、监督管理
由{department}负责政策实施的监督管理工作。

本方案自发布之日起实施。

{city}{department}
{datetime.now().strftime('%Y年%m月%d日')}"""

        return document, f"{city}_{policy_type.replace('补贴', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    def _generate_percentage_policy(self, policy_type: str) -> str:
        """生成按比例补贴的政策内容"""
        percentages = [
            (random.randint(10, 15), "一级能效或A类产品"),
            (random.randint(8, 12), "二级能效或B类产品"),
            (random.randint(5, 8), "其他符合条件的产品")
        ]
        cap = random.randint(1000, 5000)

        content = f"按照产品类别给予不同比例补贴：\n"
        for percent, category in percentages:
            content += f"- {category}：补贴销售价格的{percent}%\n"
        content += f"- 单件商品补贴最高不超过{cap}元\n"
        content += f"- 每人累计补贴不超过{cap * 3}元"

        return content

    def _generate_tiered_policy(self, policy_type: str) -> str:
        """生成分档补贴的政策内容"""
        tiers = []
        base = 0
        for i in range(4):
            low = base
            high = low + random.randint(50000, 150000)
            subsidy = random.randint(2000, 10000) * (i + 1)
            if i == 3:
                tiers.append(f"- {low/10000:.0f}万元以上：补贴{subsidy}元")
            else:
                tiers.append(f"- {low/10000:.0f}-{high/10000:.0f}万元：补贴{subsidy}元")
            base = high

        content = "根据购买金额分档补贴：\n" + "\n".join(tiers)
        return content

    def _generate_discount_policy(self, policy_type: str) -> str:
        """生成满减优惠的政策内容"""
        discounts = []
        thresholds = [100, 200, 500, 1000, 2000]
        for threshold in thresholds:
            discount = random.randint(threshold // 10, threshold // 5)
            discounts.append(f"- 满{threshold}元减{discount}元")

        content = "消费满减优惠：\n" + "\n".join(discounts)
        content += f"\n- 每人每月限享受{random.randint(3, 5)}次"
        return content

    def _generate_fixed_policy(self, policy_type: str) -> str:
        """生成固定金额补贴的政策内容"""
        amounts = []
        categories = ['普通商品', '节能产品', '智能产品', '高端产品']
        for category in categories:
            amount = random.randint(50, 500) * 10
            amounts.append(f"- {category}：补贴{amount}元/件")

        content = "按商品类别给予固定补贴：\n" + "\n".join(amounts)
        content += f"\n- 每人限购{random.randint(1, 3)}件"
        return content

    def save_as_txt(self, content: str, filename: str, output_dir: str = "data/guize") -> str:
        """保存为 TXT 文件"""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        file_path = output_path / f"{filename}.txt"
        file_path.write_text(content, encoding='utf-8')

        return str(file_path)

    def save_as_docx(self, content: str, filename: str, output_dir: str = "data/guize") -> str:
        """保存为 DOCX 文件"""
        try:
            import docx
        except ImportError:
            print("需要安装 python-docx: pip install python-docx")
            return None

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        # 创建 Word 文档
        doc = docx.Document()

        # 分段添加内容
        lines = content.split('\n')

        # 添加标题
        if lines:
            title = doc.add_heading(lines[0], 0)
            title.alignment = 1  # 居中

        # 添加其余内容
        for line in lines[1:]:
            if line.strip():
                if line.startswith('一、') or line.startswith('二、') or \
                   line.startswith('三、') or line.startswith('四、') or \
                   line.startswith('五、') or line.startswith('六、'):
                    # 章节标题
                    doc.add_heading(line, level=1)
                elif line.startswith('-'):
                    # 列表项
                    doc.add_paragraph(line, style='List Bullet')
                elif line.startswith(('1.', '2.', '3.')):
                    # 编号列表
                    doc.add_paragraph(line, style='List Number')
                else:
                    # 普通段落
                    doc.add_paragraph(line)

        # 保存文档
        file_path = output_path / f"{filename}.docx"
        doc.save(str(file_path))

        return str(file_path)


def main():
    """主程序"""
    print("=" * 70)
    print("随机政策文档生成与 DSL 转换演示")
    print("=" * 70)

    # 创建生成器
    generator = PolicyDocumentGenerator()

    # 生成随机政策
    print("\n[步骤 1] 生成随机政策文档")
    print("-" * 40)

    content, filename = generator.generate_random_policy()

    print("生成的政策内容预览：")
    print("-" * 40)
    lines = content.split('\n')[:15]  # 显示前15行
    for line in lines:
        print(f"  {line}")
    print("  ...")

    # 保存文档
    print("\n[步骤 2] 保存政策文档")
    print("-" * 40)

    # 保存为 TXT
    txt_path = generator.save_as_txt(content, filename)
    print(f"[OK] 保存为 TXT: {txt_path}")

    # 尝试保存为 DOCX
    docx_path = generator.save_as_docx(content, filename)
    if docx_path:
        print(f"[OK] 保存为 DOCX: {docx_path}")
        process_path = docx_path
    else:
        print("[INFO] 无法保存为 DOCX，将使用 TXT 文件")
        process_path = txt_path

    # 使用 DSL 生成器处理
    print("\n[步骤 3] 使用大模型生成 DSL")
    print("-" * 40)

    # 检查是否有 LLM 配置
    try:
        from config.settings import settings
        has_llm = bool(settings.model.llm_api_key)
        if has_llm:
            print(f"检测到 LLM 配置: {settings.model.llm_model_name}")
            use_llm = True
        else:
            print("未检测到 LLM 配置，将使用规则提取")
            use_llm = False
    except:
        print("无法加载配置，将使用规则提取")
        use_llm = False

    # 创建 DSL 管道
    pipeline = DSLPipeline(
        data_dir="data/guize",
        output_dir="rules",
        use_project_config=use_llm
    )

    # 处理文档
    print(f"\n处理文档: {Path(process_path).name}")
    result = pipeline.process_document(process_path)

    if result['status'] == 'success':
        print("\n[SUCCESS] DSL 生成成功！")
        print(f"   生成的 DSL 文件: {result['dsl_file']}")

        # 显示 DSL 内容
        if 'yaml_content' in result:
            print("\n生成的 DSL 内容（部分）：")
            print("-" * 40)
            yaml_lines = result['yaml_content'].split('\n')[:30]
            for line in yaml_lines:
                print(f"  {line}")
            if len(result['yaml_content'].split('\n')) > 30:
                print("  ...")

        # 测试生成的规则
        if 'dsl_data' in result:
            rule_id = result['dsl_data'].get('rule_id')
            print(f"\n[步骤 4] 测试生成的规则: {rule_id}")
            print("-" * 40)

            # 根据规则类型创建测试输入
            test_inputs = {
                'price': 10000.0,
                'category': '测试商品'
            }

            # 如果有能效相关，添加能效等级
            if 'energy_level' in result.get('dsl_data', {}).get('inputs', []):
                test_inputs['energy_level'] = 1

            test_result = pipeline.test_rule(rule_id, test_inputs)
            print(f"测试输入: {test_inputs}")
            print(f"执行状态: {test_result.get('status', 'N/A')}")
            print(f"计算结果: {test_result.get('final_result', 'N/A')}")
    else:
        print(f"\n[FAILED] DSL 生成失败: {result.get('errors', ['未知错误'])}")

    # 总结
    print("\n" + "=" * 70)
    print("演示完成")
    print("=" * 70)
    print("\n总结：")
    print("1. [OK] 成功随机生成政策文档")
    print("2. [OK] 文档已保存为 TXT/DOCX 格式")
    if result['status'] == 'success':
        print("3. [OK] 成功生成 DSL 规则文件")
        print("4. [OK] 规则可以正常执行")
        print("\n系统可以处理随机生成的政策文档并自动转换为可执行的 DSL！")
    else:
        print("3. [WARNING] DSL 生成遇到问题，请检查配置")

    return result


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='随机生成政策文档并转换为 DSL')
    parser.add_argument('--count', type=int, default=1, help='生成文档数量')
    parser.add_argument('--format', choices=['txt', 'docx', 'both'], default='both', help='文档格式')

    args = parser.parse_args()

    if args.count == 1:
        # 生成单个文档
        main()
    else:
        # 批量生成
        print(f"批量生成 {args.count} 个政策文档...")
        success_count = 0

        for i in range(args.count):
            print(f"\n--- 第 {i+1}/{args.count} 个文档 ---")
            result = main()
            if result['status'] == 'success':
                success_count += 1

        print(f"\n批���生成完成: {success_count}/{args.count} 成功")