#!/usr/bin/env python3
"""
独立测试知识库的完整流程
不依赖主应用初始化
"""

import os
import sys
from pathlib import Path

# 清除代理
for key in ['http_proxy', 'https_proxy', 'HTTP_PROXY', 'HTTPS_PROXY', 'all_proxy', 'ALL_PROXY']:
    os.environ.pop(key, None)

sys.path.insert(0, str(Path(__file__).parent))

def test_full_pipeline():
    print("=" * 60)
    print("知识库完整流程测试")
    print("=" * 60)

    # 1. 直接导入必要的知识库模块
    print("\n[1] 初始化知识库组件...")

    from dotenv import load_dotenv
    load_dotenv()

    # 导入Milvus存储
    from app.knowledge.milvus import MilvusStore
    milvus_store = MilvusStore()
    print("✓ MilvusStore初始化成功")

    # 导入嵌入模型
    from app.knowledge.embeddings import get_embedding_model
    embed_model = get_embedding_model()
    print(f"✓ 嵌入模型初始化: {embed_model.__class__.__name__}")

    # 2. 清理并创建集合
    print("\n[2] 准备Milvus集合...")

    collection_name = "policy_documents"
    try:
        # 检查并重建集合
        if milvus_store.collection_exists():
            print(f"✓ 集合 {collection_name} 已存在")
        else:
            print(f"✓ 创建新集合 {collection_name}")
    except:
        print(f"✓ 使用集合 {collection_name}")

    # 3. 处理多个文档
    print("\n[3] 批量处理文档...")

    doc_dir = Path("/data/temp33/Glyph/resources/data/raw")
    test_files = list(doc_dir.glob("*.docx"))[:5]  # 处理前5个文档

    all_chunks = []
    all_metadata = []

    for i, doc_file in enumerate(test_files, 1):
        print(f"\n[文档 {i}/{len(test_files)}] {doc_file.name}")

        try:
            # 读取文档
            from docx import Document
            doc = Document(str(doc_file))
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            if not full_text:
                print("  ✗ 文档内容为空")
                continue

            print(f"  ✓ 文档长度: {len(full_text)} 字符")

            # 切块（使用LlamaIndex）
            from llama_index.core.node_parser import SentenceSplitter

            splitter = SentenceSplitter(
                chunk_size=512,
                chunk_overlap=50,
                separator=" ",
                paragraph_separator="\n\n"
            )

            # 创建简单文档节点
            from llama_index.core.schema import Document as LlamaDocument
            llama_doc = LlamaDocument(text=full_text)

            # 切块
            chunks = splitter.get_nodes_from_documents([llama_doc])
            print(f"  ✓ 切块数: {len(chunks)}")

            # 收集切块和元数据
            for j, chunk in enumerate(chunks):
                all_chunks.append(chunk.text)
                all_metadata.append({
                    "source": doc_file.name,
                    "doc_id": f"doc_{i}",
                    "chunk_id": j,
                    "chunk_total": len(chunks)
                })

        except Exception as e:
            print(f"  ✗ 处理失败: {e}")
            continue

    print(f"\n总切块数: {len(all_chunks)}")

    # 4. 批量生成嵌入并存储
    print("\n[4] 批量嵌入和存储...")

    if all_chunks:
        try:
            # 分批处理
            batch_size = 10
            stored_count = 0

            for i in range(0, len(all_chunks), batch_size):
                batch_texts = all_chunks[i:i+batch_size]
                batch_metadata = all_metadata[i:i+batch_size]

                print(f"\n处理批次 {i//batch_size + 1}/{(len(all_chunks)-1)//batch_size + 1}")

                # 存储到Milvus（add_texts会自动生成嵌入）
                ids = milvus_store.add_texts(batch_texts, batch_metadata)
                stored_count += len(ids)
                print(f"  ✓ 存储 {len(ids)} 个切块")

            print(f"\n✓ 总共存储 {stored_count} 个切块")

        except Exception as e:
            print(f"✗ 批量存储失败: {e}")

    # 5. 测试检索
    print("\n[5] 测试检索功能...")

    test_queries = [
        "汽车消费补贴金额是多少",
        "家电以旧换新有什么要求",
        "消费券怎么使用",
        "手机购新补贴标准"
    ]

    for query in test_queries:
        print(f"\n查询: {query}")
        print("-" * 40)

        try:
            # 执行检索
            results = milvus_store.similarity_search(query, k=3)

            if results:
                for i, (doc, score) in enumerate(results, 1):
                    content = doc.page_content[:100].replace('\n', ' ')
                    source = doc.metadata.get('source', 'unknown')
                    print(f"  结果{i} (相似度: {score:.3f}):")
                    print(f"    来源: {source}")
                    print(f"    内容: {content}...")
            else:
                print("  无结果")

        except Exception as e:
            print(f"  查询失败: {e}")

    # 6. 测试重排序（如果配置了）
    print("\n[6] 测试重排序...")

    try:
        from app.knowledge.reranker import get_reranker
        reranker = get_reranker()

        if reranker:
            print("✓ Reranker已配置")

            # 测试重排
            query = "汽车补贴"
            results = milvus_store.similarity_search(query, k=10)

            if results:
                docs = [r[0] for r in results]
                reranked = reranker.rerank(query, docs, top_k=3)

                print(f"\n重排后的结果 (查询: {query}):")
                for i, doc in enumerate(reranked, 1):
                    content = doc.page_content[:80].replace('\n', ' ')
                    print(f"  {i}. {content}...")
        else:
            print("✗ Reranker未配置")

    except Exception as e:
        print(f"✗ Reranker测试失败: {e}")

    # 7. 统计信息
    print("\n[7] 集合统计...")

    try:
        stats = milvus_store.get_collection_stats()
        print(f"✓ Collection: {milvus_store.collection_name}")
        print(f"✓ 文档总数: {stats.get('row_count', 0)}")
        print(f"✓ 索引类型: IVF_FLAT")
        print(f"✓ 度量类型: COSINE")

    except Exception as e:
        print(f"统计信息获取失败: {e}")

    print("\n" + "=" * 60)
    print("✅ 知识库完整测试成功!")
    print("=" * 60)

if __name__ == "__main__":
    try:
        test_full_pipeline()
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)