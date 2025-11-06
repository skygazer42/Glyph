"""
文档解析模块
支持 .docx 和 .txt 格式的文档解析
"""

import re
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
import json

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器，支持多种格式"""

    def __init__(self):
        self.supported_formats = ['.docx', '.txt', '.doc', '.pdf']

    def parse(self, file_path: str) -> str:
        """
        解析文档文件并返回文本内容

        Args:
            file_path: 文档文件路径

        Returns:
            提取的文本内��
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {suffix}")

        if suffix == '.txt':
            return self._parse_txt(path)
        elif suffix in ['.docx', '.doc']:
            return self._parse_docx(path)
        elif suffix == '.pdf':
            return self._parse_pdf(path)
        else:
            raise ValueError(f"暂不支持的文件格式: {suffix}")

    def _parse_txt(self, path: Path) -> str:
        """解析纯文本文件"""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()

    def _parse_docx(self, path: Path) -> str:
        """解析 Word 文档"""
        try:
            import docx
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = docx.Document(str(path))
        text_parts = []

        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    def _parse_pdf(self, path: Path) -> str:
        """解析 PDF 文档"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2")

        text_parts = []
        with open(path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text.strip())

        return "\n".join(text_parts)

    def preprocess_text(self, text: str) -> Dict[str, Any]:
        """
        预处理文本，提取结构化信息

        Args:
            text: 原始文本

        Returns:
            预处理后的结构化信息
        """
        processed = {
            'raw_text': text,
            'sections': [],
            'metadata': {}
        }

        # 提取文档标题
        title_pattern = r'[《]([^》]+)[》]|^(.{1,50}(?:实施细则|管理办法|通知|公告|方案))'
        title_match = re.search(title_pattern, text[:500], re.MULTILINE)
        if title_match:
            processed['metadata']['title'] = title_match.group(1) or title_match.group(2)

        # 提取文档编号
        doc_id_pattern = r'[A-Za-z0-9\u4e00-\u9fa5]+[〔【\[][\d]+[〕】\]][\d]+号'
        doc_id_match = re.search(doc_id_pattern, text[:500])
        if doc_id_match:
            processed['metadata']['doc_id'] = doc_id_match.group()

        # 提取时间信息
        date_patterns = [
            r'(\d{4})[年-](\d{1,2})[月-](\d{1,2})[日]?',
            r'自(\d{4}[年-]\d{1,2}[月-]\d{1,2}[日]?)起',
            r'至(\d{4}[年-]\d{1,2}[月-]\d{1,2}[日]?)止'
        ]

        dates = []
        for pattern in date_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                dates.append(match.group())

        if dates:
            processed['metadata']['dates'] = dates

        # 分段处理
        sections = self._split_sections(text)
        processed['sections'] = sections

        # 提取金额信息
        amount_pattern = r'(\d+(?:\.\d+)?)[元万千百]'
        amounts = re.findall(amount_pattern, text)
        if amounts:
            processed['metadata']['amounts'] = amounts

        # 提取比例信息
        percentage_pattern = r'(\d+(?:\.\d+)?)[%％]'
        percentages = re.findall(percentage_pattern, text)
        if percentages:
            processed['metadata']['percentages'] = percentages

        return processed

    def _split_sections(self, text: str) -> List[Dict[str, str]]:
        """
        将文本分割成章节

        Args:
            text: 原始文本

        Returns:
            章节列表
        """
        sections = []

        # 按条款分割
        clause_pattern = r'第[一二三四五六七八九十\d]+[条章节]'
        clauses = re.split(clause_pattern, text)
        clause_headers = re.findall(clause_pattern, text)

        if len(clause_headers) > 0:
            # 添加第一部分（如果有）
            if clauses[0].strip():
                sections.append({
                    'header': '前言',
                    'content': clauses[0].strip()
                })

            # 添加各条款
            for i, header in enumerate(clause_headers):
                if i + 1 < len(clauses):
                    content = clauses[i + 1].strip()
                    if content:
                        sections.append({
                            'header': header,
                            'content': content
                        })
        else:
            # 没有明显的条款结构，按段落分割
            paragraphs = text.split('\n\n')
            for i, para in enumerate(paragraphs):
                if para.strip():
                    sections.append({
                        'header': f'段落{i+1}',
                        'content': para.strip()
                    })

        return sections

    def extract_tables(self, text: str) -> List[List[List[str]]]:
        """
        提取文本中的表格数据

        Args:
            text: 包含表格的文本

        Returns:
            表格数据列表
        """
        tables = []

        # 查找表格模式（使用|分隔的行）
        lines = text.split('\n')
        current_table = []

        for line in lines:
            if '|' in line:
                # 分割单元格
                cells = [cell.strip() for cell in line.split('|')]
                # 过滤空单元格
                cells = [cell for cell in cells if cell]
                if cells:
                    current_table.append(cells)
            else:
                # 非表格行，保存当前表格（如果有）
                if current_table and len(current_table) > 1:
                    tables.append(current_table)
                current_table = []

        # 保存最后一个表格
        if current_table and len(current_table) > 1:
            tables.append(current_table)

        return tables

    def extract_rules(self, text: str) -> List[Dict[str, Any]]:
        """
        提取文本中的规则信息

        Args:
            text: 包含规则的文本

        Returns:
            规则信息列表
        """
        rules = []

        # 提取满减规则
        discount_pattern = r'满(\d+(?:\.\d+)?)[元]?[减送](\d+(?:\.\d+)?)[元]?'
        discount_matches = re.finditer(discount_pattern, text)
        for match in discount_matches:
            rules.append({
                'type': 'discount',
                'condition': float(match.group(1)),
                'benefit': float(match.group(2))
            })

        # 提取阶梯规则
        tier_pattern = r'(\d+(?:\.\d+)?)[万元]?[以]?[上下]'
        tier_matches = re.finditer(tier_pattern, text)
        tiers = []
        for match in tier_matches:
            tiers.append(float(match.group(1)))

        if tiers:
            rules.append({
                'type': 'tier',
                'thresholds': sorted(set(tiers))
            })

        # 提取比例补贴规则
        subsidy_pattern = r'补贴[比例率]?[为是]?(\d+(?:\.\d+)?)[%％]'
        subsidy_matches = re.finditer(subsidy_pattern, text)
        for match in subsidy_matches:
            rules.append({
                'type': 'percentage',
                'rate': float(match.group(1)) / 100
            })

        return rules